from flask import Blueprint, jsonify, request, send_file
import logging
import io
import time
import traceback
from docx import Document
from docx.shared import Inches, Pt 
from docx.enum.text import WD_ALIGN_PARAGRAPH

import mariadb # <--- AÑADIDO para mariadb.Error
import mysql.connector # <--- AÑADIDO por si se usa mysql.connector.Error

# Configuración del logging
logger = logging.getLogger(__name__)

# Intentar importar módulos necesarios
try:
    from controlador.arp_manager import start_full_attack, stop_full_attack, get_mac_for_ipv4, mac_to_ipv6_linklocal, GATEWAY_IPV6_LL, initialize_network_info as initialize_arp_network_info
    from controlador.iptables_manager import block_ipv4, unblock_ipv4, block_ipv6, unblock_ipv6
    from modelo.base_datos import obtener_conexion 
    ATTACK_MODULES_AVAILABLE = True
    logger.info("✅ Módulos de ataque, iptables y DB cargados para acciones de alerta.")
except ImportError as e:
    logger.error(f"⚠️ ADVERTENCIA en alertas_acciones: No se pudo importar un módulo esencial ({e}). Funcionalidades de bloqueo/ataque estarán limitadas.")
    ATTACK_MODULES_AVAILABLE = False
    
    def _dummy_attack_func(ip_addr, *args, **kwargs): logger.warning(f"Función de ataque no disponible, llamada dummy para IP: {ip_addr}"); return False
    def _dummy_mac_func(ip_addr): logger.warning(f"Función get_mac no disponible, llamada dummy para IP: {ip_addr}"); return None
    def _dummy_linklocal_func(mac_addr): logger.warning(f"Función mac_to_ipv6_linklocal no disponible, llamada dummy para MAC: {mac_addr}"); return None
    
    start_full_attack = _dummy_attack_func
    stop_full_attack = _dummy_attack_func
    get_mac_for_ipv4 = _dummy_mac_func
    mac_to_ipv6_linklocal = _dummy_linklocal_func
    GATEWAY_IPV6_LL = None
    initialize_arp_network_info = lambda: False 

    block_ipv4 = _dummy_attack_func
    unblock_ipv4 = _dummy_attack_func
    block_ipv6 = _dummy_attack_func
    unblock_ipv6 = _dummy_attack_func
    
    try:
        from modelo.base_datos import obtener_conexion
    except ImportError:
        logger.critical("FALLO CRÍTICO: No se pudo importar obtener_conexion de modelo.base_datos.")
        obtener_conexion = lambda: None 


alertas_acciones_bp = Blueprint('alertas_acciones', __name__, url_prefix='/alerta')

# --- NUEVA FUNCIÓN ---
def obtener_estado_actual_dispositivo(target_ip):
    """
    Consulta el estado 'bloqueado' de un dispositivo en la BD.
    Retorna:
        1 si está bloqueado.
        0 si no está bloqueado.
        None si la IP no se encuentra o hay un error de BD.
    """
    if obtener_conexion is None:
        logger.error(f"DB_GET_STATUS: La función obtener_conexion no está disponible. No se puede obtener estado para {target_ip}.")
        return None

    conexion = None
    cursor = None
    estado_bloqueado = None # Por defecto None para indicar error o no encontrado

    try:
        conexion = obtener_conexion()
        if not conexion:
            logger.error(f"DB_GET_STATUS: No se pudo obtener conexión a la BD para IP {target_ip}.")
            return None # Error de conexión

        cursor = conexion.cursor()
        sql_select = "SELECT bloqueado FROM dispositivos WHERE direccion_ip = %s"
        cursor.execute(sql_select, (target_ip,))
        resultado = cursor.fetchone()

        if resultado:
            estado_bloqueado = int(resultado[0]) # Convertir a entero (0 o 1)
            logger.info(f"DB_GET_STATUS: Estado 'bloqueado' para IP {target_ip} es {estado_bloqueado}.")
        else:
            logger.warning(f"DB_GET_STATUS: IP {target_ip} no encontrada en la tabla dispositivos.")
            # Mantenemos estado_bloqueado = None para indicar que no se encontró
            # O podría devolverse un valor específico si se prefiere distinguir "no encontrado" de "error DB"
            
    except (mariadb.Error, mysql.connector.Error) as db_err:
        logger.error(f"DB_GET_STATUS: Error de base de datos al obtener estado para {target_ip}: {db_err}")
        estado_bloqueado = None # Error de base de datos
    except Exception as e:
        logger.error(f"DB_GET_STATUS: Error inesperado al obtener estado para {target_ip}: {e}")
        logger.error(traceback.format_exc())
        estado_bloqueado = None # Otro error
    finally:
        if cursor:
            try:
                cursor.close()
            except Exception as e_cur:
                logger.error(f"DB_GET_STATUS: Error cerrando cursor para {target_ip}: {e_cur}")
        if conexion and conexion.is_connected():
            try:
                conexion.close()
            except Exception as e_conn:
                logger.error(f"DB_GET_STATUS: Error cerrando conexión para {target_ip}: {e_conn}")
    
    return estado_bloqueado

# --- FUNCIÓN EXISTENTE (MODIFICADA) ---
def actualizar_estado_bloqueado_db(target_ip, bloqueado_status_int):
    """
    Actualiza el estado 'bloqueado' (0 o 1) en la tabla dispositivos.
    Retorna True si fue exitoso y al menos una fila fue afectada o el estado ya era el deseado, False en caso contrario.
    """
    if obtener_conexion is None: 
        logger.error(f"DB_UPDATE_BLOCKED: La función obtener_conexion no está disponible. No se puede actualizar {target_ip}.")
        return False

    conexion = None
    cursor = None
    operacion_exitosa = False 
    try:
        conexion = obtener_conexion()
        if not conexion:
            logger.error(f"DB_UPDATE_BLOCKED: No se pudo obtener conexión a la BD para actualizar {target_ip}.")
            return False

        cursor = conexion.cursor() 

        logger.info(f"DB_UPDATE_BLOCKED: Intentando actualizar 'bloqueado' a {bloqueado_status_int} para IP {target_ip}")
        sql_update = "UPDATE dispositivos SET bloqueado = %s WHERE direccion_ip = %s"
        cursor.execute(sql_update, (bloqueado_status_int, target_ip))
        
        conexion.commit()
        logger.info(f"DB_UPDATE_BLOCKED: Commit realizado. Estado de {target_ip} actualizado a {bloqueado_status_int}. Filas afectadas: {cursor.rowcount}")
        
        if cursor.rowcount > 0:
            operacion_exitosa = True
        else:
            # Si no afectó filas, verificar si es porque ya estaba en ese estado
            estado_actual_post_intento = obtener_estado_actual_dispositivo(target_ip)
            if estado_actual_post_intento == bloqueado_status_int:
                logger.info(f"DB_UPDATE_BLOCKED: 0 filas afectadas, pero el estado para {target_ip} ya es {bloqueado_status_int}.")
                operacion_exitosa = True # Considerar éxito si ya estaba en el estado deseado
            else:
                logger.warning(f"DB_UPDATE_BLOCKED: 0 filas afectadas al actualizar {target_ip} y el estado no es el deseado. ¿Existe la IP?")
                # operacion_exitosa sigue False

    except (mariadb.Error, mysql.connector.Error) as db_err: 
        logger.error(f"DB_UPDATE_BLOCKED: Error de base de datos al actualizar {target_ip}: {db_err}")
        if conexion:
            try:
                conexion.rollback()
                logger.info(f"DB_UPDATE_BLOCKED: Rollback realizado para {target_ip} debido a error.")
            except Exception as e_rb:
                logger.error(f"DB_UPDATE_BLOCKED: Error durante el rollback para {target_ip}: {e_rb}")
    except Exception as e:
        logger.error(f"DB_UPDATE_BLOCKED: Error inesperado al actualizar {target_ip}: {e}")
        logger.error(traceback.format_exc()) 
    finally:
        if cursor:
            try:
                cursor.close()
            except Exception as e_cur:
                logger.error(f"DB_UPDATE_BLOCKED: Error cerrando cursor para {target_ip}: {e_cur}")
        if conexion and conexion.is_connected(): 
            try:
                conexion.close()
            except Exception as e_conn:
                logger.error(f"DB_UPDATE_BLOCKED: Error cerrando conexión para {target_ip}: {e_conn}")
    return operacion_exitosa


@alertas_acciones_bp.route('/bloquear_ip/<path:target_ip>', methods=['POST'])
def bloquear_ip_desde_alerta(target_ip):
    if not target_ip or not target_ip.strip() or target_ip.strip().lower() in ['null', 'none', 'undefined']:
        logger.error(f"BLOQUEO_IP_VALIDATION: IP inválida o nula recibida: '{target_ip}'")
        return jsonify({"status": "error", "mensaje": f"❌ IP inválida: '{target_ip}'. No se puede bloquear."}), 400

    logger.info(f"BLOQUEO_IP_VALIDATION: Solicitud de bloqueo para IP: {target_ip}")

    if not ATTACK_MODULES_AVAILABLE:
        logger.warning("BLOQUEO_IP_VALIDATION: Módulos de ataque no disponibles.")
        return jsonify({"status": "error", "mensaje": "❌ Funcionalidad de ataque/bloqueo no disponible en el servidor."}), 503

    # --- INICIO DE LA NUEVA VALIDACIÓN ---
    estado_actual = obtener_estado_actual_dispositivo(target_ip)

    if estado_actual == 1:
        mensaje_ya_bloqueado = f"ℹ️ El dispositivo {target_ip} ya se encuentra bloqueado."
        logger.info(f"BLOQUEO_IP_VALIDATION: {mensaje_ya_bloqueado}")
        return jsonify({"status": "info", "mensaje": mensaje_ya_bloqueado}), 200 # 200 OK con mensaje informativo
    
    if estado_actual is None:
        # La IP no está en la tabla 'dispositivos'. Puede que quieras bloquearla igualmente
        # a nivel de red y luego decidir si la añades a la tabla o no.
        # Por ahora, se registrará la advertencia desde obtener_estado_actual_dispositivo
        # y el flujo continuará para bloquear a nivel de red.
        # La actualización de DB (actualizar_estado_bloqueado_db) fallará en afectar filas
        # si la IP no existe, lo cual ya se maneja.
        logger.warning(f"BLOQUEO_IP_VALIDATION: IP {target_ip} no encontrada en la tabla dispositivos. Se intentará bloqueo de red.")
    # --- FIN DE LA NUEVA VALIDACIÓN ---


    logger.info(f"BLOQUEO_IP_EXEC: Procediendo con el intento de bloqueo para IP: {target_ip} (Estado actual en BD: {estado_actual})")

    if not initialize_arp_network_info(): 
        logger.error("BLOQUEO_IP_EXEC: Fallo al inicializar la información de red para arp_manager.")
        # Continuar con el bloqueo de iptables si es posible.

    ipv4_blocked_iptables = block_ipv4(target_ip)
    ipv6_blocked_iptables = False 
    mitm_iniciado = start_full_attack(target_ip) # Esta función debería tener su propia lógica para no sobre-escribir o causar conflicto si ya está activo.
                                                 # Sin embargo, nuestra validación previa ya debería evitar llamarla si está bloqueado en BD.
    if not mitm_iniciado:
        logger.warning(f"BLOQUEO_IP_EXEC: Fallo al iniciar MiTM para {target_ip}, pero se intentará bloqueo con iptables.")

    target_mac_for_ipv6 = get_mac_for_ipv4(target_ip) 
    if target_mac_for_ipv6 and GATEWAY_IPV6_LL:
        target_ipv6_ll = mac_to_ipv6_linklocal(target_mac_for_ipv6)
        if target_ipv6_ll:
            logger.info(f"BLOQUEO_IP_EXEC: Intentando bloqueo IPv6 con iptables para {target_ipv6_ll} (derivado de {target_ip})")
            ipv6_blocked_iptables = block_ipv6(target_ipv6_ll)
        else:
            logger.warning(f"BLOQUEO_IP_EXEC: No se pudo derivar IPv6 LL para MAC {target_mac_for_ipv6} de IP {target_ip}")
    else:
        logger.warning(f"BLOQUEO_IP_EXEC: No se bloqueará IPv6 con iptables para {target_ip} (Falta MAC o Gateway IPv6_LL).")

    db_updated = False 
    if ipv4_blocked_iptables or ipv6_blocked_iptables or mitm_iniciado:
        # Solo actualizamos si alguna acción de red fue (aparentemente) exitosa
        logger.info(f"BLOQUEO_IP_EXEC: Al menos una acción de bloqueo de red para {target_ip} tuvo éxito. Intentando actualizar BD.")
        db_updated = actualizar_estado_bloqueado_db(target_ip, 1) 
        if not db_updated:
            logger.error(f"BLOQUEO_IP_EXEC: El bloqueo de red para {target_ip} pudo haber funcionado, pero falló la actualización en BD o la IP no existe/ya estaba actualizada.")
    else:
        logger.error(f"BLOQUEO_IP_EXEC: Fallaron todos los métodos de bloqueo (iptables y MiTM) para {target_ip}.")
        # Si las acciones de red fallaron, no tiene sentido devolver 200
        return jsonify({"status": "error", "mensaje": f"❌ Falló el bloqueo completo a nivel de red para {target_ip}."}), 500

    mensajes_exito = []
    if mitm_iniciado: mensajes_exito.append("MiTM iniciado")
    if ipv4_blocked_iptables: mensajes_exito.append("IPv4 bloqueado (iptables)")
    if ipv6_blocked_iptables: mensajes_exito.append("IPv6 bloqueado (iptables)")
    
    final_msg_parts = []
    if mensajes_exito:
        final_msg_parts.append(f"Acciones de red para {target_ip}: {', '.join(mensajes_exito)}.")
    else:
        # Esto no debería pasar si llegamos aquí porque la condición anterior (if ipv4_blocked_iptables or ...)
        # habría resultado en un error 500. Pero por si acaso:
        final_msg_parts.append(f"Ninguna acción de bloqueo de red tuvo éxito explícito para {target_ip}.")

    if db_updated:
        final_msg_parts.append("Estado en BD actualizado a bloqueado.")
    else:
        # Si llegamos aquí, significa que alguna acción de red tuvo éxito, pero la BD no se actualizó
        # (o la IP no existe en la tabla, o ya estaba marcada como bloqueada, lo cual la validación inicial ya debería haber capturado).
        # La función actualizar_estado_bloqueado_db ya loggea los detalles.
        final_msg_parts.append("El estado en BD no se actualizó (podría no existir la IP en tabla o ya estar marcada).")
        
    final_msg = " ".join(final_msg_parts)
    logger.info(f"BLOQUEO_IP_EXEC: Resultado final: {final_msg}")
    
    # Si al menos una acción de red tuvo éxito, consideramos la operación general como exitosa.
    if mensajes_exito:
        return jsonify({"status": "success", "mensaje": final_msg }), 200
    else:
        # Este caso es redundante debido al chequeo previo, pero por seguridad:
        return jsonify({"status": "error", "mensaje": final_msg}), 500


@alertas_acciones_bp.route('/generar_informe', methods=['POST'])
def generar_informe_alerta_docx():
    try:
        alerta_data = request.json
        if not alerta_data:
            return jsonify({"status": "error", "mensaje": "No se recibieron datos de la alerta."}), 400

        document = Document()
        
        document.add_heading('Informe de Alerta de Seguridad', level=1)
        p = document.add_paragraph()
        p.add_run('Fecha del Informe: ').bold = True
        p.add_run(time.strftime("%Y-%m-%d %H:%M:%S"))

        def add_data_row(label, value):
            p_row = document.add_paragraph()
            p_row.add_run(f"{label}: ").bold = True
            p_row.add_run(str(value) if value is not None else 'N/A')

        add_data_row('ID Alerta', alerta_data.get('id_evento_seguridad', alerta_data.get('id_alerta', 'N/A')))
        add_data_row('Tipo de Alerta', alerta_data.get('tipo'))
        add_data_row('Nivel de Severidad', alerta_data.get('nivel'))
        add_data_row('Fecha de Detección', alerta_data.get('fecha')) 
        add_data_row('Repeticiones', alerta_data.get('repeticiones'))
        
        document.add_heading('Descripción de la Amenaza', level=2)
        document.add_paragraph(str(alerta_data.get('descripcion', 'N/A')))

        document.add_heading('Detalles de Red Involucrados', level=2)
        add_data_row('IP Origen', alerta_data.get('ip_origen'))
        add_data_row('MAC Origen', alerta_data.get('mac_origen'))
        add_data_row('Puerto Origen', alerta_data.get('puerto_origen'))
        add_data_row('SO Origen (Estimado)', alerta_data.get('so_origen'))
        
        add_data_row('IP Destino', alerta_data.get('ip_destino'))
        add_data_row('MAC Destino', alerta_data.get('mac_destino'))
        add_data_row('Puerto Destino', alerta_data.get('puerto_destino'))
        add_data_row('Protocolo', alerta_data.get('protocolo'))

        document.add_heading('Estado del Evento (en el IDS)', level=2)
        add_data_row('Estado de la Alerta', alerta_data.get('estado_alerta'))

        if alerta_data.get('detalles'):
            document.add_heading('Detalles Adicionales del Paquete', level=2)
            document.add_paragraph(str(alerta_data.get('detalles')))
            
        file_stream = io.BytesIO()
        document.save(file_stream)
        file_stream.seek(0)

        tipo_alerta_seguro = "".join(c if c.isalnum() else "_" for c in alerta_data.get('tipo', 'generico'))
        filename = f"informe_alerta_{tipo_alerta_seguro}_{time.strftime('%Y%m%d_%H%M%S')}.docx"
        
        logger.info(f"INFORME_DOCX: Generando informe: {filename}")
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        logger.error(f"INFORME_DOCX: Error al generar informe: {e}")
        logger.error(traceback.format_exc())
        return jsonify({"status": "error", "mensaje": "Error interno al generar el informe."}), 500